# coding: utf-8
"""
Calculates VFR routes where the legs are defined as a function
"""
# general packages
import json
from pathlib import Path
from typing import Optional, Union
import matplotlib.axes
from typing_extensions import Self
import textwrap
from enum import Enum, auto
import datetime
import os
import io
import math
from contextlib import contextmanager

import requests

# projection related packages
from pyproj import Proj, Geod
import numpy as np

# pdf and imaging related packages
import pymupdf
import PIL
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.patches as patches

# document creation related packages
from docx import Document
from docx.shared import Cm

# gpx read and create
from lxml import etree # pylint: disable=no-name-in-module
import gpxpy

# LaTeX evaluation related imports
from sympy.utilities.lambdify import lambdify
import sympy.abc

# package imports
from .projutils import (
    PointLonLat, PointXY,
    ExtentLonLat, ExtentXY,
    _calculate_2d_transformation_matrix,
    _apply_transformation_matrix,
    _rotate_point,
    _get_extent_from_points,
    _get_extent_from_extents,
    parse_latex_with_constants
)
from .docxutils import add_formula_par

OPENWEATHER_ENDPOINT = "https://api.openweathermap.org/data/2.5/forecast?lat={lat}&lon={lon}&appid={OPENWEATHER_APIKEY}"
MAGDEV_ENDPOINT = "https://www.ngdc.noaa.gov/geomag-web/calculators/calculateDeclination?lat1={lat}&lon1={lon}&startYear={when.year}&startMonth={when.month}&startDay={when.day}&resultFormat=json&key={MAGDEV_APIKEY}"
OPENWEATHER_APIKEY = os.getenv("OPENWEATHER_APIKEY")
MAGDEV_APIKEY = os.getenv("MAGDEV_APIKEY")

class VFRRouteState(Enum):
    INITIATED = auto()
    AREAOFINTEREST = auto()
    WAYPOINTS = auto()
    LEGS = auto()
    ANNOTATIONS = auto()
    FINALIZED = auto()
    
class VFRCoordSystem(Enum):
    FUNCTION = auto()
    MAPCROP_XY = auto()
    MAP_XY = auto()
    FULL_WORLD_XY = auto()
    LONLAT = auto()


class VFRPoint:
    """
    A Point object which knows its coordinate system and coordinates.
    Holds neccessary references to be able to transform from one system to another.
    """
    def __init__(self,
                 x: float, y: float,
                 coord_system: Optional[VFRCoordSystem] = VFRCoordSystem.LONLAT,
                 route: Optional["VFRRoute"] = None,
                 leg: Optional["VFRLeg"] = None):
        """
        Initialize the point
        Args
            coord_system: VFRCoordSystem
                The coordinate system this points coordinates are in
            x: float
                The horizontal coordinate or longitude (if it is a global point)
            y: float
                The vertical coordinate or latitude (if it is a global point)
        """
        self.coord_system = coord_system
        self.x, self.y = x, y
        self.route = route
        self.leg = leg
            
    @property
    def lon(self):
        return self.x
        
    @property
    def lat(self):
        return self.y

    def toDict(self):
        """
        Converts this object to a JSON serializable dictionary.
        WARNING: References to `route` and `leg` are NOT preserved!
        """
        return {
            'x': self.x,
            'y': self.y,
            'coord_system': self.coord_system.name
        }
    

    @classmethod
    def fromDict(cls, value, route: Union['VFRFunctionRoute', None] = None, leg: Union['VFRLeg', None] = None):
        """
        Converts a dictionary from toDict into a VFRPoint.
        WARNING: since references were not saved they can be passed to
        this method.
        """
        return VFRPoint(value['x'], value['y'], VFRCoordSystem[value['coord_system']], route, leg)


    def project_point(self, to_system: VFRCoordSystem) -> "VFRPoint":
        """
        Project this point to another coordinate system.
        It uses the parameters in the referenced route and leg (if neccessary)
        """
        if (not self.leg) and (self.coord_system==VFRCoordSystem.FUNCTION or to_system==VFRCoordSystem.FUNCTION):
            raise ValueError("There is no leg reference defined and you tried to convert to/from function coordinate system.")
        if self.coord_system==to_system:
            return self
        curx, cury, cursys = self.x, self.y, self.coord_system
        if self.coord_system.value > to_system.value:
            if cursys == VFRCoordSystem.LONLAT and cursys.value > to_system.value:
                (curx, cury), cursys = self.route._proj(curx, cury), VFRCoordSystem.FULL_WORLD_XY
            if cursys == VFRCoordSystem.FULL_WORLD_XY and cursys.value > to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.route._matrix_fullmap2map), VFRCoordSystem.MAP_XY
            if cursys == VFRCoordSystem.MAP_XY and cursys.value > to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.route._matrix_map2cropmap), VFRCoordSystem.MAPCROP_XY
            if cursys == VFRCoordSystem.MAPCROP_XY and cursys.value > to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.leg._matrix_cropmap2func), VFRCoordSystem.FUNCTION
            return VFRPoint(curx, cury, cursys, self.route, self)
        else:
            if cursys == VFRCoordSystem.FUNCTION and cursys.value < to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.leg._matrix_func2cropmap), VFRCoordSystem.MAPCROP_XY
            if cursys == VFRCoordSystem.MAPCROP_XY and cursys.value < to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.route._matrix_cropmap2map), VFRCoordSystem.MAP_XY
            if cursys == VFRCoordSystem.MAP_XY and cursys.value < to_system.value:
                (curx, cury), cursys = _apply_transformation_matrix((curx, cury), self.route._matrix_map2fullmap), VFRCoordSystem.FULL_WORLD_XY
            if cursys == VFRCoordSystem.FULL_WORLD_XY and cursys.value < to_system.value:
                (curx, cury), cursys = self.route._proj(curx, cury, inverse=True), VFRCoordSystem.LONLAT
            return VFRPoint(curx, cury, cursys, self.route, self)



class VFRAnnotation:
    """
    """

    USE_SAMPLE_WEATHER = os.getenv('USE_SAMPLE_WEATHER', "True").lower() in ["true", "yes", "on", "1"]
    BACKGROUND_COLOR = (1.0, 0.7, 0.7, 0.99)

    def __init__(self, 
                 leg: "VFRLeg",
                 name: str,
                 x: float,
                 ofs: tuple[float, float]):
        """
        """
        self._leg: VFRLeg = leg
        self.name = name
        self.x = x
        self.ofs = ofs
        self._seglen = None
        self._seglens = None
        self._segtime = None
        self._times_withwind = None
        self._weather = None
        self._headings = None
        self._declination = None


    def _clear_cache(self):
        # clear cache of cached items
        self._seglen = None
        self._seglens = None
        self._segtime = None
        self._times_withwind = None
        self._weather = None
        self._headings = None
        self._declination = None
        

    def __repr__(self):
        """
        """
        return f"{type(self).__name__}({self.name}, {self.x})"
    

    def toDict(self):
        return {
            'name': self.name,
            'x': self.x,
            'ofs': self.ofs
        }
    
    @classmethod
    def fromDict(cls, value, leg: 'VFRLeg'):
        return VFRAnnotation(leg, value['name'], value['x'], value['ofs'])
    
    
    @property
    def seglen(self):
        if self._seglen:
            return self._seglen
        x0, x1 = self._leg.ann_start_end(self)
        # calc segment points
        x = np.linspace(x0, x1, 100)
        psrc = [VFRPoint(x, self._leg.function(x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg) for x in x]
        ps = [p.project_point(VFRCoordSystem.LONLAT) for p in psrc]
        # calc segment length
        self._seglen = self._leg._route._geod.line_length([p.lon for p in ps], [p.lat for p in ps])
        return self._seglen
    

    @property
    def seglens(self):
        if self._seglens:
            return self._seglens
        x0, x1 = self._leg.ann_start_end(self)
        # calc segment points
        x = np.linspace(x0, x1, 100)
        psrc = [VFRPoint(x, self._leg.function(x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg) for x in x]
        ps = [p.project_point(VFRCoordSystem.LONLAT) for p in psrc]
        # calc segment length
        self._seglens = self._leg._route._geod.line_lengths([p.lon for p in ps], [p.lat for p in ps])
        return self._seglens


    @property
    def segtime(self):
        if self._segtime:
            return self._segtime
        seglen = self.seglen
        self._segtime = seglen/1852/self._leg._route.speed*60 if seglen else None
        return self._segtime
    

    @property
    def times_withwind(self):
        if self._times_withwind:
            return self._times_withwind
        headings = [h if h>=0 else h+360 for h in self.headings]
        wind_corrections = self.wind_corrections()
        speeds_withwind = [(self._leg._route.speed*math.cos(math.radians(-1*wind_corrections[i])))+(self.wind_speed*math.cos(math.radians(headings[i]+wind_corrections[i]-self.wind_dir+180))) for i in range(len(headings))]
                          #(speed*COS(RADIÁN(-wind_correction)))+(wind_speed*COS(RADIÁN(heading+wind_correction-wind_direction+180)))
        self._times_withwind = [self.seglens[i]/1852/speeds_withwind[i]*60 for i in range(len(self.seglens))]
        return self._times_withwind



    @property
    def headings(self) -> list[float]:
        if self._headings:
            return self._headings
        x0, x1 = self._leg.ann_start_end(self)
        x = np.linspace(x0, x1, 100)
        psrc = [VFRPoint(x, self._leg.function(x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg) for x in x]
        ps = [p.project_point(VFRCoordSystem.LONLAT) for p in psrc]
        lat = [p.lat for p in ps]
        lon = [p.lon for p in ps]
        headings, _, _ = self._leg._route._geod.inv(lon[:-1], lat[:-1], lon[1:], lat[1:])
        def clamp(deg):
            while deg>360: deg-=360
            while deg<0: deg+=360
            return deg
        headings = [clamp(h) for h in headings]
        self._headings = headings
        return self._headings


    def get_weather(self):
        # only download once
        if not self._weather:
            # download weather at from point
            if not self.USE_SAMPLE_WEATHER:
                p = VFRPoint(self.x, self._leg.function(self.x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg)
                p = p.project_point(VFRCoordSystem.LONLAT)
                response = self._leg._route._session.get(OPENWEATHER_ENDPOINT.format(
                    lon=p.lon,
                    lat=p.lat,
                    OPENWEATHER_APIKEY=OPENWEATHER_APIKEY
                ))
                self._weather = response.json()
            else:
                import json
                with open(os.path.join(Path(__file__).parent, 'sample_weather.json'), 'rt', encoding='utf8') as f:
                    self._weather = json.load(f)


    def magnetic_deviation(self, when=datetime.datetime.now()):
        if self._declination:
            return self._declination
        try:
            p = VFRPoint(self.x, self._leg.function(self.x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg)
            p = p.project_point(VFRCoordSystem.LONLAT)
            api_res = self._leg._route._session.get(MAGDEV_ENDPOINT.format(
                lon = p.lon,
                lat = p.lat,
                MAGDEV_APIKEY = MAGDEV_APIKEY,
                when = when
            ))
            api_res = api_res.json()
            self._declination = api_res["result"][0]["declination"]
        except:
            import traceback
            traceback.print_exc()
            self._declination = 5
        return self._declination

    
    @property
    def wind(self):
        self.get_weather()
        weather_ts = int(self._leg._route.dof.timestamp())
        return sorted((wfx for wfx in self._weather['list'] if wfx['dt']<=weather_ts), key=lambda wfx: wfx['dt'])[-1]['wind']


    @property
    def wind_speed(self) -> float:
        return self.wind['speed']*3600/1852


    @property
    def wind_dir(self) -> float:
        return self.wind['deg']


    def wind_corrections(self, speed: Optional[float] = None, headings: Optional[list[float]] = None) -> list[float]:
        if not speed:
            speed = self._leg._route.speed
        if not headings:
            headings = self.headings
        return [math.degrees(math.asin(self.wind_speed/self._leg._route.speed*math.sin(math.radians(h-self.wind_dir+180)))) for h in headings]
               #FOK(ARCSIN(wind_speed/speed*SIN(RADIÁN(heading-wind_direction+180))))



    def draw(self, ax: matplotlib.axes.Axes):
        xy = VFRPoint(self.x, self._leg.function(self.x), VFRCoordSystem.FUNCTION, self._leg._route, self._leg).project_point(VFRCoordSystem.MAPCROP_XY)
        seglen = self.seglen
        segtime = self.segtime
        segtime_wind = sum(self.times_withwind)
        seghdgs = self.headings
        seghdg = seghdgs[-1]
        wind_corrs = self.wind_corrections(headings=seghdgs)
        wind_corr = wind_corrs[-1] if wind_corrs else None
        mag_dev = self.magnetic_deviation(self._leg._route.dof)
        s_seglen = f"\nhossz: {seglen/1852:.1f}NM\nidő: {math.floor(segtime):3d}:{math.floor((segtime-math.floor(segtime))*60):02d}" + \
                   f" / {math.floor(segtime_wind):3d}:{math.floor((segtime_wind-math.floor(segtime_wind))*60):02d}"
        if self._leg.annotations.index(self) == 0:
            s_seglen = ""
        ax.annotate(f'{self.name}\nirány: ${self.headings[-1]:.0f}\\degree${mag_dev:+.0f}(M){wind_corr:+.0f}(W:{self.wind_speed:.0f}/{self.wind_dir:.0f}){s_seglen}',
                    xy=(xy.x, xy.y), xycoords='data',
                    xytext=(self.ofs[0], self.ofs[1]), textcoords='offset points',
                    size=5.5, va="center",
                    bbox=dict(boxstyle="round", fc=self.BACKGROUND_COLOR, ec="none"),
                    arrowprops=dict(arrowstyle="wedge,tail_width=1.",
                                    fc=self.BACKGROUND_COLOR, ec="none",
                                    patchA=None,
                                    patchB=None,
                                    relpos=(0.2, 0.5)))


class VFRLeg:
    """
    """
    def __init__(self,
                 route: "VFRRoute",
                 name: str, 
                 function_name: str,
                 function_range: str,
                 function,
                 points: list[tuple[VFRPoint, float]]):
        """
        """
        self._route: "VFRRoute" = route
        self.name = name
        self.function_name = function_name
        self.function_range = function_range
        self.function = function
        self.points = points
        for p, x in self.points:
            p.leg = self
        self.annotations: list[VFRAnnotation] = []

        self.color="red"
        self.lw=2
        
        
    def add_annotation(self, name: str, x: float, ofs: tuple[float, float]) -> Self:
        """
        """
        self._route._ensure_state(VFRRouteState.LEGS)
        newannotation = VFRAnnotation(self, name, x, ofs)
        self.annotations.append(newannotation)
        return self
        
        
    def get_extent(self) -> ExtentLonLat:
        return _get_extent_from_points([PointLonLat(p.lon, p.lat) for p, x in self.points])
        
        
    def draw(self, ax, with_annotations: bool = True):
        # draw planned track
        x = np.linspace(min([x for p, x in self.points]), 
                        max([x for p, x in self.points]),
                        100
                       )
        psrc = [VFRPoint(x, self.function(x), VFRCoordSystem.FUNCTION, self._route, self) for x in x]
        ps = [p.project_point(VFRCoordSystem.MAPCROP_XY) for p in psrc]
        ax.plot([p.x for p in ps],
                [p.y for p in ps],
                color=self.color,
                lw=self.lw
               )
        # draw annotations
        if with_annotations:
            for a in self.annotations:
                a.draw(ax)
    
    
    def calc_function(self):
        try:
            parsedfun = parse_latex_with_constants(self.function_name) # parse_latex(latex)
            self.function = lambdify(sympy.abc.x, parsedfun, modules=["math"])
        except Exception as e:
            self.function = lambda x: x # fallback to linear


    def calc_transformations(self):
        self.calc_function()
        sp = [(x, self.function(x)) for (p, x) in self.points]
        dpp = [p.project_point(VFRCoordSystem.MAPCROP_XY) for p, x in self.points]
        dp = [(p.x, p.y) for p in dpp]
        if len(sp) < 3:
            sp.append(_rotate_point(sp[0], sp[1], -90))
            dp.append(_rotate_point(dp[0], dp[1], 90))
        try:
            self._matrix_func2cropmap = _calculate_2d_transformation_matrix(sp, dp)
            self._matrix_cropmap2func = np.linalg.inv(self._matrix_func2cropmap)
        except:
            pass # keep the old matrix
        
        
    def transform_point(self,
                        p: tuple[float, float],
                        source: VFRCoordSystem,
                        dest: VFRCoordSystem
                       ) -> tuple[float, float]:
        return self._route.transform_point(p, source, dest, self)
    

    def ann_start_end(self, ann: VFRAnnotation) -> float:
        # calc start and end x
        i = self.annotations.index(ann)
        x1 = self.annotations[i].x
        if i > 0:
            x0 = self.annotations[i-1].x
        else:
            dirmul = -1.0 if self.annotations[i].x < self.annotations[i+1].x else 1.0
            x0 = x1 + dirmul*0.00001
        return x0, x1


    def toDict(self):
        return {
            'name': self.name,
            'function_name': self.function_name,
            'function_range': self.function_range,
            'points': [{'p': p.toDict(), 'x': x} for p, x in self.points]
        }
    
    @classmethod
    def fromDict(cls, value, route: Union['VFRFunctionRoute', None]):
        return VFRLeg(route,
                      value['name'],
                      value['function_name'],
                      value['function_range'],
                      None,
                      [(VFRPoint.fromDict(pdef['p'], route), pdef['x']) for pdef in value['points']])
    
    def __repr__(self):
        """
        """
        s = f"{type(self).__name__}({self.name}, {self.function_name})\n"
        for a in self.annotations:
            s += textwrap.indent(repr(a), "  ")+"\n"
        return s
        
        
class VFRTrack:
    """
    """
    def __init__(self, route: "VFRRoute", fname: Union[str, Path], color: str, xmlb: Optional[bytes] = None, load: bool = True):
        self._route = route
        self.fname = fname
        self.color = color
        self.points: list[VFRPoint] = []
        if load:
            self.points=self.read_gpx(fname=fname, xmlb=xmlb)
        
    def read_gpx(self, fname: Union[str, Path] = None, xmlb: Optional[bytes] = None) -> list[VFRPoint]:
        """Reads flown points from a GPX file / GPX string"""
        if xmlb:
            plangpx = etree.fromstring(xmlb)
        else:
            plangpx = etree.parse(fname)
        ns = {"gpx": "http://www.topografix.com/GPX/1/1", "geotracker": "http://ilyabogdanovich.com/gpx/extensions/geotracker"}
        planptx = plangpx.xpath("/gpx:gpx/gpx:trk/gpx:trkseg/gpx:trkpt", namespaces=ns)
        planpts: list[VFRPoint] = []
        for i, ptx in enumerate(planptx):
            lon = float(ptx.get('lon'))
            lat = float(ptx.get('lat'))
            planpts.append(
                VFRPoint(lon, lat, VFRCoordSystem.LONLAT, self._route)
            )
            prevlon = lon
            prevlat = lat
        return planpts
    
    def draw(self, ax: matplotlib.axes.Axes):
        # draw track
        ps = [p.project_point(VFRCoordSystem.MAPCROP_XY) for p in self.points]
        ax.plot([p.x for p in ps],
                [p.y for p in ps],
                color=self.color,
                lw=2
               )
        
    def toDict(self):
        return {
            'name': self.fname,
            'color': self.color,
            'points': [p.toDict() for p in self.points]
        }

    @classmethod
    def fromDict(cls, value, route: 'VFRFunctionRoute'):
        trk = VFRTrack(route, value['name'], value['color'], load=False)
        trk.points = [VFRPoint.fromDict(p, route) for p in value['points']]
        return trk
        
    def get_extent(self):
        return _get_extent_from_points([PointLonLat(p.lon, p.lat) for p in self.points])


class VFRFunctionRoute:
    """
    A class that can be used to
      - gradually build up a route
      - add flight tracks
      - generate a map of the route and tracks
      - generate a flight plan document
      - generate a flight plan file for SkyDaemon
    """
    
    # "https://ais.hungarocontrol.hu/aip/2024-03-21-ICAO500k/ICAO500_2024_egylapos.pdf"
    PDF_URL = os.getenv("PDF_URL")
    PDF_FILE = os.getenv("PDF_FILE")
    PDF_IN_WORLD_XY = {
      PointLonLat(16.0, 48.5): PointXY(6.0, 54.0),
      PointLonLat(16.5, 46.0): PointXY(155.83, 1639.36),
      PointLonLat(21.0, 46.0): PointXY(2133.00, 1658.6),
    } # I don't see any possibility to do this automatically
    PDF_MARGINS = ((79, 110.1), (79.05, 139)) #(('left', 'top'), ('right', 'bottom'))
    DPI = int(os.getenv("HIGH_DPI", "600"))
    LOWDPI = int(os.getenv("LOW_DPI", "72"))

    
    def __init__(self,
                 name: str,
                 speed: float,
                 dof: datetime.datetime,
                 session: requests.Session = None,
                 workfolder: Union[str, Path, None] = None,
                 outfolder: Union[str, Path, None] = None,
                 tracksfolder: Union[str, Path, None] = None
                ):
        """
        """
        self._state = VFRRouteState.INITIATED
        self.name = name
        self.speed = speed
        self.dof = dof
        self.workfolder = workfolder
        self.outfolder = outfolder
        self.tracksfolder = tracksfolder
        self.legs: list[VFRLeg] = []
        self.tracks: list[VFRTrack] = []
        self._session = session if session else requests.Session()
        self._lowresmap = None
        self._basemapimg = None
        self.waypoints: list[tuple[str, VFRPoint]] = []
        self.area_of_interest = {
            'top-left': VFRPoint(18.5, 47.5, VFRCoordSystem.LONLAT, self),
            'bottom-right': VFRPoint(19.5, 47.0, VFRCoordSystem.LONLAT, self)
        }
        self.fig, self.ax = None, None
        self.download_map()
        self.calc_extents()
        self.calc_transformations()
        

    def _ensure_state(self, required_state: VFRRouteState, ensure_minimum: bool = True, ensure_exactly: bool = False):
        """
        Ensure the object is in the desired state. Otherwise raise an exception.

        Args
            required_state: VFRRouteState
                The state we compare to
            
            ensure_minimum: bool
                The direction we compare: if True we need to have at least
                the given state, if False we need to have at most the given
                state.

            ensure_exactly: bool
                We have to have exactly the given state (no more, no less).
        """
        if ensure_exactly:
            if self._state.value!=required_state.value:
                raise RuntimeError(f"VFRFunctionRoutes object not in required state: Current {self._state}, required exact state: {required_state}.")
        elif ensure_minimum:
            if self._state.value<required_state.value:
                raise RuntimeError(f"VFRFunctionRoutes object not in required state: Current {self._state}, required minimum state: {required_state}.")
        else:
            if self._state.value>required_state.value:
                raise RuntimeError(f"VFRFunctionRoutes object not in required state: Current {self._state}, required maximum state: {required_state}.")
        
    
    def set_state(self, required_state: VFRRouteState):
        if self._state==required_state:
            return
        if self._state.value<required_state.value: # forward stepping
            if self._state == VFRRouteState.INITIATED and required_state.value > self._state.value:
                # INITIADED -> AREAOFINTEREST
                self._basemapimg = None
                self._state = VFRRouteState.AREAOFINTEREST
                self.calc_extents()
                self.calc_transformations()
            if self._state==VFRRouteState.AREAOFINTEREST and required_state.value>self._state.value:
                # AREAOFINTEREST -> WAYPOINTS
                self._state = VFRRouteState.WAYPOINTS
                self.waypoints_to_legs()
                self.calc_extents()
                self.calc_transformations()
            if self._state == VFRRouteState.WAYPOINTS and required_state.value > self._state.value:
                # WAYPOINTS -> LEGS
                self._state = VFRRouteState.LEGS
                self.legs_to_annotations()
                self.calc_extents()
                self.calc_transformations()
            if self._state == VFRRouteState.LEGS and required_state.value > self._state.value:
                # LEGS -> ANNOTATIONS
                self._state = VFRRouteState.ANNOTATIONS
            if self._state == VFRRouteState.ANNOTATIONS and required_state.value > self._state.value:
                # ANNOTATIONS -> FINALIZED
                self.finalize()
        else: # backward stepping
            # TODO: currently nothing is done (we could free up resources, etc)
            self._state = required_state


    def waypoints_to_legs(self):
        """
        Converts waypoints to legs considering the already existing ones
        """
        for i, wp_start in enumerate(self.waypoints):
            wp_end = self.waypoints[i+1 if i+1<len(self.waypoints) else 0] # circle around (last point is the same as first)
            if len(self.legs)>i: # we have a leg at that position
                leg = self.legs[i]
                # so we adjust its endpoints position (not the x value)
                leg.points[0] = (VFRPoint(wp_start[1].lon, wp_start[1].lat, VFRCoordSystem.LONLAT, self), leg.points[0][1])
                leg.points[-1] = (VFRPoint(wp_end[1].lon, wp_end[1].lat, VFRCoordSystem.LONLAT, self), leg.points[-1][1])
                # we adjust the name of the leg
                leg.name = f"{wp_start[0]} -- {wp_end[0]}"
                # we adjust the annotations so we have the first and last match
                if len(leg.annotations)>0:
                    leg.annotations[0].x = leg.points[0][1]
                else:
                    leg.add_annotation('???', leg.points[0][1], (0,0))
                if len(leg.annotations)>1:
                    leg.annotations[-1].x = leg.points[-1][1]
                else:
                    leg.add_annotation('???', leg.points[-1][1], (0,0))
            else: # we don't have a leg yet
                # so we add a new one
                self.add_leg(f"{wp_start[0]} -- {wp_end[0]}", f"x^{i+1}", "", lambda x: x**i, 
                             [
                                 (VFRPoint(wp_start[1].lon, wp_start[1].lat, VFRCoordSystem.LONLAT, self), 0),
                                 (VFRPoint(wp_end[1].lon, wp_end[1].lat, VFRCoordSystem.LONLAT, self), 1)
                             ])


    def legs_to_annotations(self):
        """
        Transfers information from leg points to annotations (at least a start and an end is needed)
        """
        for leg in self.legs:
            if len(leg.annotations) > 0:
                leg.annotations[0].x = leg.points[0][1]
            else:
                leg.add_annotation('???', leg.points[0][1], (0, 0))
            if len(leg.annotations) > 1:
                leg.annotations[-1].x = leg.points[-1][1]
            else:
                leg.add_annotation('???', leg.points[-1][1], (0, 0))


    def finalize(self):
        self._ensure_state(VFRRouteState.ANNOTATIONS)
        self.calc_extents()
        self.calc_transformations()
        self.calc_basemap()
        # TODO: obtain live data (from internet)
        self._state = VFRRouteState.FINALIZED
        
        
    def download_map(self):
        """Downloads map as pdf from internet if not already exists as a file."""
        self.pdf_destination = os.path.join(self.workfolder, type(self).PDF_FILE)
        if not os.path.isfile(self.pdf_destination):
            response = self._session.get(type(self).PDF_URL)
            with open(self.pdf_destination, 'wb') as pdf_file:
                pdf_file.write(response.content)


    @contextmanager
    def get_lowres_map(self):
        # state check
        self._ensure_state(VFRRouteState.INITIATED)

        # pdf conversion and caching
        if not self._lowresmap:
            pdf_document = pymupdf.open(self.pdf_destination)
            page = pdf_document[0]
            rect = page.rect  # the page rectangle
            m = self.PDF_MARGINS
            clip = pymupdf.Rect(m[0][0], m[0][1], rect.width-m[1][0], rect.height-m[1][1])  # the area we want
            pdfimage_cropcheck = page.get_pixmap(clip=clip, dpi=self.LOWDPI)
            pilimage_cropcheck = PIL.Image.open(io.BytesIO(pdfimage_cropcheck.tobytes("png")))
            self._lowresmap = pilimage_cropcheck

        # draw
        fig = plt.figure()
        fig.set_size_inches((c/self.LOWDPI for c in self._lowresmap.size))
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)
        ax.imshow(self._lowresmap)

        # return
        yield fig

        # TODO: cleanup
        plt.close(fig)


    def set_area_of_interest(self, top_left_x: float, top_left_y: float, bottom_right_x: float, bottom_right_y: float) -> None:
        self._ensure_state(VFRRouteState.INITIATED)
        self.area_of_interest = {
            'top-left': VFRPoint(top_left_x, top_left_y, VFRCoordSystem.MAP_XY, self).project_point(VFRCoordSystem.LONLAT),
            'bottom-right': VFRPoint(bottom_right_x, bottom_right_y, VFRCoordSystem.MAP_XY, self).project_point(VFRCoordSystem.LONLAT)
        }

    @contextmanager
    def get_highres_map(self):
        # state check
        self._ensure_state(VFRRouteState.AREAOFINTEREST)

        # pdf conversion and caching
        if not self._basemapimg:
            self.calc_basemap()

        # draw
        fig = plt.figure()
        fig.set_size_inches((c/self.DPI for c in self._basemapimg.size))
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)
        ax.imshow(self._basemapimg)

        # return
        yield fig, ax

        # TODO: cleanup
        plt.close(fig)


    @contextmanager
    def get_annotations_map(self):
        with self.get_highres_map() as (fig, ax):
            for l in self.legs:
                l.draw(ax, False)
            yield fig, ax            


    @contextmanager
    def get_tracks_map(self):
        with self.get_highres_map() as (fig, ax):
            for l in self.legs:
                l.draw(ax)
            for t in self.tracks:
                t.draw(ax)
            yield fig, ax            


    def add_waypoint(self, name: str, point: VFRPoint):
        point.route = self
        self.waypoints.append((name, point.project_point(VFRCoordSystem.LONLAT)))


    def update_waypoints(self, wps: list[dict]):
        # calculate new waypoints
        self.waypoints = [(wp["name"], VFRPoint(wp["x"], wp["y"], VFRCoordSystem.MAPCROP_XY, self).project_point(VFRCoordSystem.LONLAT)) for wp in wps]


    def update_legs(self, legs: list[dict]):
        # set legs according to edits
        for i, leg in enumerate(legs):
            curleg = self.legs[i]
            # setup general
            curleg.name = leg["name"]
            # setup function
            curleg.function_range = leg["function_range"]
            latex = leg["function_name"]
            curleg.function_name = latex
            curleg.calc_function()
            # setup constraint points
            lp_start, lp_end = curleg.points[0], curleg.points[-1]
            newpoints: list[tuple[VFRPoint, float]] = []
            for j, pt in enumerate(leg["points"]):
                if j==0:
                    newpoints.append((lp_start[0], pt["func_x"]))
                elif j==len(leg["points"])-1:
                    newpoints.append((lp_end[0], pt["func_x"]))
                else:
                    newpoints.append((
                        VFRPoint(pt["x"], pt["y"], VFRCoordSystem.MAPCROP_XY, self, curleg).project_point(VFRCoordSystem.LONLAT),
                        pt["func_x"]
                    ))
            curleg.points = newpoints
            # recalculate
            curleg.calc_transformations()


    def update_annotations(self, legs: list[dict]):
        for i, l in enumerate(legs):
            if i>len(self.legs)-1:
                break
            if self.legs[i].name!=l['name']:
                print(f"WARNING: leg number {i} name does not match ({self.legs[i].name}!={l['name']})")
            self.legs[i].annotations = [VFRAnnotation(self.legs[i],
                                                      a['name'],
                                                      a['func_x'],
                                                      (a['ofs']['x'], a['ofs']['y']))
                                        for a in l['annotations']
                                       ]


    def _fullmap_clicker(self):
        """
        Allows to get click points on the map (to later estimate transformations from them)
        """
        pdf_document = pymupdf.open(self.pdf_destination)
        page = pdf_document[0]
        rect = page.rect  # the page rectangle
        m = self.PDF_MARGINS
        clip = pymupdf.Rect(m[0][0], m[0][1], rect.width-m[1][0], rect.height-m[1][1])  # the area we want
        pdfimage_cropcheck = page.get_pixmap(clip=clip)
        pilimage_cropcheck = PIL.Image.open(io.BytesIO(pdfimage_cropcheck.tobytes("png")))
        fig, ax = plt.subplots()
        print(pilimage_cropcheck.size)
        ax.imshow(pilimage_cropcheck)
        # draw reference points
        p = [self._proj(ll.lon, ll.lat) for ll in self.PDF_IN_WORLD_XY.keys()]
        p = [_apply_transformation_matrix(pp, self._matrix_fullmap2map) for pp in p]
        ax.scatter([pp[0] for pp in p], [pp[1] for pp in p], marker='X', c='red')
        # draw map rectangle
        p = [
            PointLonLat(self.extent.minlon, self.extent.minlat), # minlon-minlat -> leftbottom
            PointLonLat(self.extent.minlon, self.extent.maxlat), # minlon-maxlat -> lefttop
            PointLonLat(self.extent.maxlon, self.extent.minlat), # maxlon-minlat -> rightbottom
            PointLonLat(self.extent.maxlon, self.extent.maxlat), # maxlon-maxlat -> righttop
        ]
        p = [self._proj(ll.lon, ll.lat) for ll in p]
        p = [_apply_transformation_matrix(pp, self._matrix_fullmap2map) for pp in p]
        ax.scatter([pp[0] for pp in p], [pp[1] for pp in p], marker='X', c='red')
        p = self.get_mapxyextent()
        rect = patches.Rectangle((p.minx, p.miny), p.maxx-p.minx, p.maxy-p.miny, linewidth=1, edgecolor='r', facecolor='none')
        ax.add_patch(rect)
        # interactive plot
        from matplotlib.backend_bases import MouseButton
        def on_click(event):
            if event.button is MouseButton.LEFT:
                print(event)
        plt.connect('button_press_event', on_click)
        plt.show()


    def _map_clicker(self):
        """
        Allows to get click points on the map (to later estimate transformations from them)
        """
        xt = self.get_mapxyextent()
        x0, y0, x1, y1 = xt.minx, xt.miny, xt.maxx, xt.maxy
        scale = self.LOWDPI/self.DPI #1 / (1/72*2.54/100*500000)
        pdf_document = pymupdf.open(self.pdf_destination)
        page = pdf_document[0]
        rect = page.rect  # the page rectangle
        m = self.PDF_MARGINS
        clip = pymupdf.Rect(m[0][0]+x0,
                            m[0][1]+y0,
                            m[0][0]+x1,  # rect.width-m[1][0]
                            m[0][1]+y1,  # rect.height-m[1][1]
                           )
        print(rect, clip)
        pdfimage_cropcheck = page.get_pixmap(clip=clip, dpi=self.DPI)
        pilimage_cropcheck = PIL.Image.open(io.BytesIO(pdfimage_cropcheck.tobytes("png")))
        fig, ax = plt.subplots()
        print(pilimage_cropcheck.size)
        ax.imshow(pilimage_cropcheck)
        # draw map rectangle
        print("points:")
        p = [
            PointLonLat(self.extent.minlon, self.extent.minlat), # minlon-minlat -> leftbottom
            PointLonLat(self.extent.minlon, self.extent.maxlat), # minlon-maxlat -> lefttop
            PointLonLat(self.extent.maxlon, self.extent.minlat), # maxlon-minlat -> rightbottom
            PointLonLat(self.extent.maxlon, self.extent.maxlat), # maxlon-maxlat -> righttop
        ]
        p = [self._proj(ll.lon, ll.lat) for ll in p]
        p = [_apply_transformation_matrix(pp, self._matrix_fullmap2map) for pp in p]
        print(f"  fullmap: {p}")
        p = [_apply_transformation_matrix(pp, self._matrix_map2cropmap) for pp in p]
        print(f"  cropmap: {p}")
        ax.scatter([pp[0] for pp in p], [pp[1] for pp in p], marker='X', c='red')
        print("rect:")
        p = [(xt.minx, xt.miny), (xt.maxx, xt.maxy)]
        print(f" fullmap: {p}")
        p = [_apply_transformation_matrix(pp, self._matrix_map2cropmap) for pp in p]
        print(f" cropmap: {p}")
        rect = patches.Rectangle((p[0][0], p[0][1]), p[1][0]-p[0][0], p[1][1]-p[0][1], linewidth=1, edgecolor='r', facecolor='none')
        ax.add_patch(rect)
        # draw all leg points
        ps = []
        ps2 = []
        for l in self.legs:
            for p, x in l.points:
                pp = p.project_point(VFRCoordSystem.MAPCROP_XY)
                ps.append(pp)
                pp = self._proj(p.lon, p.lat)
                pp = _apply_transformation_matrix(pp, self._matrix_fullmap2map)
                pp = _apply_transformation_matrix(pp, self._matrix_map2cropmap)
                ps2.append(pp)
        ax.scatter([pp.x for pp in ps], [pp.y for pp in ps], marker='X', c='blue')
        ax.scatter([pp[0] for pp in ps2], [pp[1] for pp in ps2], marker='X', c='green')
        # interactive plot
        from matplotlib.backend_bases import MouseButton
        def on_click(event):
            if event.button is MouseButton.LEFT:
                print(event)
        plt.connect('button_press_event', on_click)
        plt.show()

        
    def add_leg(self, 
                name: str,
                function_name: str,
                function_range: str,
                function,
                points: list[tuple[VFRPoint, float]]) -> VFRLeg:
        """
        """
        self._ensure_state(VFRRouteState.WAYPOINTS)
        for p, x in points:
            p.route = self
        newleg = VFRLeg(self, name, function_name, function_range, function, points)
        self.legs.append(newleg)
        return newleg
        
    def add_track(self,
                  fname: Union[str, Path],
                  color: str,
                  xmlb: Optional[bytes] = None):
        """Adds a flown track to show on the map.
        
        Args:
            fname: str
                The filename of the track (it will be looked for in the trackfolders folder
                if xmlstring is not given)
            color: str
        """
        self._ensure_state(VFRRouteState.ANNOTATIONS)
        self.tracks.append(VFRTrack(self, fname, color, xmlb=xmlb))
        return self
    

    def update_tracks(self, tracks):
        newtracks: list[VFRTrack] = []
        for t in self.tracks:
            if t.fname in [nt['name'] for nt in tracks]:
                newtracks.append(t)
                t.color = [nt['color'] for nt in tracks if nt['name']==t.fname][0]
        self.tracks = newtracks
        
        
    def calc_transformations(self):
        """Calculates and saves the transformations between coordinate systems.
        
        The following coordinate systems are neccessary:
            1. the function coordinate system
            2. cropped map image coordinates
            3. the map image coordinate system
            4. full-world x-y coordinates
            5. the latitude-longitude coordinates on the earth
            
        It is neccessary to transform between those coordinate systems:
            - 1->2 to draw the route on the image
            - 1->5 to save the route to a flightplan
            - 5->2 to draw tracks on the map image
        Problem:
            2 through 5 is map dependent
            1 is function (leg) dependent
        """
        # to consider:
        #   map false_easting, false_northing values calc/hardcoded
        #   map scale
        # calculate transformations for LONLAT<->FULL_WORLD_XY
        self._proj_str = "+proj=lcc +lon_0=-90 +lat_1=46 +lat_2=48 +ellps=WGS84"
        self._proj = Proj(self._proj_str)
        self._geod = Geod(self._proj_str)
        #print(self._proj(16,48.5))
        #print(self._proj(17,47.5))
        # calculate transformations for FULL_WORLD_XY<->MAP_XY
        p = [self._proj(ll.lon, ll.lat) for ll in self.PDF_IN_WORLD_XY.keys()] # convert to fullworld map coord
        self._matrix_fullmap2map = _calculate_2d_transformation_matrix(p, list(self.PDF_IN_WORLD_XY.values())) # calc matrix from fullworld map coord to map coord
        self._matrix_map2fullmap = np.linalg.inv(self._matrix_fullmap2map)
        #print("TEST lonlat to mapxy:")
        #for i, (lonlat, xy) in enumerate(self.PDF_IN_WORLD_XY.items()):
        #    print(f"  {lonlat} -> {p[i]} -> {xy} vs {_apply_transformation_matrix(p[i], self._matrix_fullmap2map)}")
        # calculate transformations for MAP_XY<->MAPCROP_XY
        p = self.get_mapxyextent()
        pp = [
            (0, 0),
            (0, (p.maxy-p.miny)*self.DPI/self.LOWDPI),
            ((p.maxx-p.minx)*self.DPI/self.LOWDPI, 0),
            ((p.maxx-p.minx)*self.DPI/self.LOWDPI, (p.maxy-p.miny)*self.DPI/self.LOWDPI),
        ]  # also scale it up!
        p = [
            (p.minx, p.miny),
            (p.minx, p.maxy),
            (p.maxx, p.miny),
            (p.maxx, p.maxy),
        ]
        self._matrix_map2cropmap = _calculate_2d_transformation_matrix(p, pp)
        self._matrix_cropmap2map = np.linalg.inv(self._matrix_map2cropmap)
        # print("TEST mapxy to cropmapxy:")
        #for i, (mapxy, cropmapxy) in enumerate(zip(p, pp)):
        #    print(f"  {mapxy} -> {cropmapxy} vs {_apply_transformation_matrix(mapxy, self._matrix_map2cropmap)}")
        # calculate transformations for each leg
        for leg in self.legs:
            leg.calc_transformations()
        
        
    def transform_point(self,
                        p: tuple[float, float],
                        source: VFRCoordSystem,
                        dest: VFRCoordSystem,
                        leg: Optional[VFRLeg] = None
                       ) -> tuple[float, float]:
        self._ensure_state(VFRRouteState.FINALIZED)
        if source==dest:
            return p
        if (source==VFRCoordSystem.FUNCTION or dest==VFRCoordSystem.FUNCTION) and not leg:
            raise RuntimeError(f"Can not transform coordinates from {source} to {dest} without leg specified")
        raise NotImplementedError("you should implement transform_point")
        
        
    def calc_extents(self, margin_x: float = .2, margin_y: Optional[float] = None):
        """Calculates and saves the extents of the neccessary map.
        
        Calculates the top-left and bottom-right coordinates of the map
        that will contain all points of the route. It also adds a margin
        in both directions given in the parameters as a percentage.

        If there are no legs nor tracks, than it defaults to an area around Budapest.
        
        Args:
            margin_x:
                Horizontal additional margin. Given in percentage of the
                distance of minimum and maximum lateral coordinates.
                Default is 10%
            margin_y:
                Optional vertical margin. Given in percentage of the
                distance of minimum and maximum longitudinal coordinates.
                Default is equal to horizontal margin.
        """
        # extent: (min_lat, max_lat, min_lon, max_lon)
        lat0, lat1, lon0, lon1 = (0, 0, 0, 0)
        if self._state==VFRRouteState.INITIATED:
            self.extent=ExtentLonLat(18.5, 47.0, 19.5, 47.5) # just a default around Budapest
        elif self._state==VFRRouteState.AREAOFINTEREST:
            lat0, lat1, lon0, lon1 = (self.area_of_interest["top-left"].lat,
                                        self.area_of_interest["bottom-right"].lat,
                                        self.area_of_interest["top-left"].lon,
                                        self.area_of_interest["bottom-right"].lon)
            self.extent = ExtentLonLat(
                min(lon0, lon1),
                min(lat0, lat1),
                max(lon0, lon1),
                max(lat0, lat1)
            )
        elif self._state in [VFRRouteState.WAYPOINTS, VFRRouteState.LEGS, VFRRouteState.ANNOTATIONS, VFRRouteState.FINALIZED]:
            if len(self.legs)==0 and len(self.tracks)==0:
                self.extent=ExtentLonLat(18.5, 47.0, 19.5, 47.5) # just a default around Budapest
                return
            extent = _get_extent_from_extents(
                    [l.get_extent() for l in self.legs] +
                    [t.get_extent() for t in self.tracks]
                )
            if margin_y is None:
                margin_y = margin_x
            extent_with_margins = ExtentLonLat(
                    extent.minlon - (extent.maxlon-extent.minlon)*margin_x,
                    extent.minlat - (extent.maxlat-extent.minlat)*margin_y,
                    extent.maxlon + (extent.maxlon-extent.minlon)*margin_x,
                    extent.maxlat + (extent.maxlat-extent.minlat)*margin_y
                )
            self.extent = extent_with_margins


    def get_mapxyextent(self) -> ExtentXY:
        p = [
            PointLonLat(self.extent.minlon, self.extent.minlat), # minlon-minlat -> leftbottom
            PointLonLat(self.extent.minlon, self.extent.maxlat), # minlon-maxlat -> lefttop
            PointLonLat(self.extent.maxlon, self.extent.minlat), # maxlon-minlat -> rightbottom
            PointLonLat(self.extent.maxlon, self.extent.maxlat), # maxlon-maxlat -> righttop
        ]
        p = [self._proj(ll.lon, ll.lat) for ll in p] # lon, lat
        p = [_apply_transformation_matrix(pp, self._matrix_fullmap2map) for pp in p]
        p = [PointXY(pp[0], pp[1]) for pp in p]
        return ExtentXY(
            min(pp.x for pp in p),
            min(pp.y for pp in p),
            max(pp.x for pp in p),
            max(pp.y for pp in p)
        )


    def draw_map(self):
        """Draws a matplotlib based map of the defined route.
        
        Args:
        
        Returns:
            The matplotlib axes of the final plot.
        """
        self._ensure_state(VFRRouteState.FINALIZED)
        
        # initialize map
        fig = plt.figure()
        fig.set_size_inches((c/self.DPI for c in self._basemapimg.size))
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)
        
        # draw the map parts
        ax.imshow(self._basemapimg)
        for l in self.legs:
            l.draw(ax)
        for t in self.tracks:
            t.draw(ax)
        
        # save and return the image
        self.fig, self.ax = fig, ax
        return fig, ax
        
        
    def calc_basemap(self):
        # calc clip coordinates
        lat0, lat1, lon0, lon1 = (0, 0, 0, 0)
        if self._state==VFRRouteState.INITIATED:
            raise RuntimeError(f"VFRFunctionRoutes object not in required state: Current {self._state}, required: {VFRRouteState.AREAOFINTEREST}.")
        elif self._state==VFRRouteState.AREAOFINTEREST:
            lat0, lat1, lon0, lon1 = (self.area_of_interest["top-left"].lat,
                                        self.area_of_interest["bottom-right"].lat,
                                        self.area_of_interest["top-left"].lon,
                                        self.area_of_interest["bottom-right"].lon)
        elif self._state==VFRRouteState.WAYPOINTS:
            raise NotImplementedError("Sorry this is not implemented yet")
        elif self._state in [VFRRouteState.LEGS, VFRRouteState.ANNOTATIONS, VFRRouteState.FINALIZED]:
            lat0, lat1, lon0, lon1 = self.extent.minlat, self.extent.maxlat, self.extent.minlon, self.extent.maxlon
        corners_lonlat = [
            VFRPoint(lon0, lat0, route=self),
            VFRPoint(lon1, lat1, route=self),
            VFRPoint(lon1, lat0, route=self),
            VFRPoint(lon0, lat1, route=self)
        ]
        corners_map = [p.project_point(to_system=VFRCoordSystem.MAP_XY) for p in corners_lonlat]
        x0 = min([p.x for p in corners_map])
        y0 = min([p.y for p in corners_map])
        x1 = max([p.x for p in corners_map])
        y1 = max([p.y for p in corners_map])
        if y1<y0:
            y0, y1 = y1, y0 # the order of them is important
        # clip the image
        pdf_document = pymupdf.open(self.pdf_destination)
        page = pdf_document[0]
        m=((79, 110.1), (79.05, 139))
        clip = pymupdf.Rect(
            m[0][0]+x0,
            m[0][1]+y0,
            m[0][0]+x1,
            m[0][1]+y1
        )  # the area we want
        pdfimage = page.get_pixmap(clip=clip, dpi=self.DPI)
        self._basemapimg = PIL.Image.open(io.BytesIO(pdfimage.tobytes("png")))
            
            
    def create_doc(self, save: bool = True) -> Union[io.BytesIO, None]:
        """
        Generates a report of the route as a Word document
        Argument: a list of a tuple
        - first element is a tuple of the leg data
            - 1st element is the name (section header)
            - 2nd element is the function with the bounds used as a text
            - 3rd element is the definition of x points to real life as a text
        - second element is a list of segment data, which is a tuple
            - 1st element is the name
            - 2nd element is the heading at that point
            - 3rd element is the length of the curve segment in meters
            - 4th element is the time used for that curve segment
        """
        self._ensure_state(VFRRouteState.FINALIZED)
        # draw map if we don't have it yet and save the image
        if not self.fig:
            self.draw_map()
        imgname = os.path.join(self.outfolder, self.name+'.png')
        self.fig.savefig(
            imgname,         
            bbox_inches='tight',
            pad_inches=0,
            dpi=self.DPI
        )
        # header and image
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        doc.add_heading('Route Plan', 0)
        doc.add_paragraph(f"Planned speed: {self.speed} KIAS.")
        doc.add_paragraph(f"Wind forecast for {self.dof:%Y-%m-%d %H:%M %Z}.")
        doc.add_picture(f'{imgname}', width=Cm(19.00))
        doc.add_page_break()

        # legs
        totdist, tottime, tottimewc = 0, 0, 0
        for leg in self.legs:
            # leg heading and definiton
            doc.add_heading(leg.name, level=1)
            doc.add_heading("Definition", level=2)
            add_formula_par(doc, f"${leg.function_name}$", style="List Bullet")
            add_formula_par(doc, leg.function_range, style="List Bullet")
            doc.add_heading("Segments", level=2)

            # leg table header
            tab = doc.add_table(rows=1, cols=8)
            tab.autofit = True
            tab.allow_autofit = True
            tab.style = "Colorful Shading Accent 1"
            for i, hdr in enumerate(["Name", "Hdg", "Mag", "WC", "Length", "Time", "Tme(WC)", "Wind"]):
                tab.rows[0].cells[i].text = hdr

            # TODO: leg table rows (per annotations)
            for ann in leg.annotations:
                seglen = ann.seglen
                segtime = ann.segtime
                segtime_wind = sum(ann.times_withwind)
                seghdgs = ann.headings
                seghdg = seghdgs[-1]
                wind_corrs = ann.wind_corrections(headings=seghdgs)
                wind_corr = wind_corrs[-1] if wind_corrs else None
                magdev = ann.magnetic_deviation(self.dof)


                row_cells = tab.add_row().cells
            
                row_cells[0].text = ann.name
                #row_cells[0].width = Cm(2)
                
                row_cells[1].text = f"{seghdg:5.0f}\N{DEGREE SIGN}"
                row_cells[2].text = f"{magdev:3.0f}\N{DEGREE SIGN}"
                row_cells[3].text = f"{wind_corr:3.0f}\N{DEGREE SIGN}"
                row_cells[4].text = f"{seglen/1852 if seglen is not None else '-':{'' if seglen is None else '.1f'}}NM"
                row_cells[5].text = f"{math.floor(segtime) if segtime is not None else '-':{'' if segtime is None else '2d'}}"+\
                                    f":{math.floor((segtime-math.floor(segtime))*60) if segtime is not None else '--':{'' if segtime is None else '02d'}}"
                row_cells[6].text = f"{math.floor(segtime_wind) if segtime_wind is not None else '-':{'' if segtime_wind is None else '2d'}}"+\
                                    f":{math.floor((segtime_wind-math.floor(segtime_wind))*60) if segtime_wind is not None else '--':{'' if segtime_wind is None else '02d'}}"
                row_cells[7].text = f"{ann.wind_dir:3d}\N{DEGREE SIGN} {ann.wind_speed:.0f}kts"

                totdist+=seglen if seglen is not None else 0
                tottime+=segtime if segtime is not None else 0
                tottimewc+=segtime_wind if segtime_wind is not None else 0

        # total distance/time
        doc.add_paragraph(f"Total: {totdist/1852:.0f} NM, " +
                          f"{math.floor(tottime):2d}:{math.floor((tottime-math.floor(tottime))*60):02d} / " +
                          f"{math.floor(tottimewc):2d}:{math.floor((tottimewc-math.floor(tottimewc))*60):02d}"
                         )


        # save it
        if save:
            docname = os.path.join(self.outfolder, self.name+'.docx')
            doc.save(docname)
            return
        else:
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf


    def save_plan(self):
        self._ensure_state(VFRRouteState.FINALIZED)
        gpx = gpxpy.gpx.GPX()
        gpx.name = "Elmebeteg VFR útvonal"
        gpx.time = datetime.datetime.now()
        rte = gpxpy.gpx.GPXRoute(name="Elmebeteg VFR útvonal")
        for leg in self.legs:
            x = np.linspace(min([x for p, x in leg.points]),
                            max([x for p, x in leg.points]),
                            100
                            )
            psrc = [VFRPoint(x, leg.function(x), VFRCoordSystem.FUNCTION, self, leg) for x in x]
            ps = [p.project_point(VFRCoordSystem.LONLAT) for p in psrc]
            pt = [gpxpy.gpx.GPXRoutePoint(p.lat, p.lon, name=leg.name if i==0 else None) for i, p in enumerate(ps)]
            rte.points.extend(pt)
        gpx.routes.append(rte)
        return gpx.to_xml()


    def __repr__(self):
        """
        """
        s = f"{type(self).__name__}({self.name})\n"
        for l in self.legs:
            s += textwrap.indent(repr(l), "  ")
        return s
    

    def toDict(self):
        # initiate json object with basic info
        jsonrte = {
            'name': self.name,
            'speed': self.speed,
            'dof': self.dof.isoformat(),
            'state': self._state.name
        }
        # step 1: area of interest
        if self._state.value>=VFRRouteState.AREAOFINTEREST.value:
            jsonrte['step1'] = { 'area_of_interest': {
                'top-left': self.area_of_interest['top-left'].toDict(),
                'bottom-right': self.area_of_interest['bottom-right'].toDict(),
            }}
        # step 2: waypoints
        if self._state.value>=VFRRouteState.WAYPOINTS.value:
            jsonrte['step2'] = { 'waypoints': [(wp[0], wp[1].toDict()) for wp in self.waypoints] }
        # step 3: legs
        if self._state.value>=VFRRouteState.LEGS.value:
            jsonrte['step3'] = { 'legs': [leg.toDict() for leg in self.legs] }
        # step 4: annotation points
        if self._state.value>=VFRRouteState.ANNOTATIONS.value:
            jsonrte['step4'] = { 'annotations': [[ann.toDict() for ann in leg.annotations] for leg in self.legs] }
        # step 5: tracks
        if self._state.value>=VFRRouteState.FINALIZED.value:
            jsonrte['step5'] = { 'tracks': [t.toDict() for t in self.tracks]}
        # return
        return jsonrte
    
    def toJSON(self):
        return json.dumps(self.toDict(), indent=2)
    

    @classmethod
    def fromJSON(cls, jsonstring: str, 
                 session: requests.Session = None,
                 workfolder: Union[str, Path, None] = None,
                 outfolder: Union[str, Path, None] = None,
                 tracksfolder: Union[str, Path, None] = None):
        # decode json
        jsonrte = json.loads(jsonstring)
        # initiate with basic info
        rte = VFRFunctionRoute(jsonrte['name'], jsonrte['speed'], datetime.datetime.fromisoformat(jsonrte['dof']),
                               session, workfolder, outfolder, tracksfolder)
        state = VFRRouteState[jsonrte['state']]
        # step 1: area of interest
        if state.value>=VFRRouteState.AREAOFINTEREST.value:
            rte.area_of_interest = {
                'top-left': VFRPoint.fromDict(jsonrte['step1']['area_of_interest']['top-left'], rte),
                'bottom-right': VFRPoint.fromDict(jsonrte['step1']['area_of_interest']['bottom-right'], rte),
            }
            rte.set_state(VFRRouteState.AREAOFINTEREST)
        # step 2: waypoints
        if state.value>=VFRRouteState.WAYPOINTS.value:
            rte.waypoints = [(name, VFRPoint.fromDict(p, rte)) for name, p in jsonrte['step2']['waypoints']]
            rte.set_state(VFRRouteState.WAYPOINTS)
        # step 3: legs
        if state.value>=VFRRouteState.LEGS.value:
            rte.legs = [VFRLeg.fromDict(leg, rte) for leg in jsonrte['step3']['legs']]
            rte.set_state(VFRRouteState.LEGS)
        # step 4: annotation points
        if state.value>=VFRRouteState.ANNOTATIONS.value:
            for i, l in enumerate(jsonrte['step4']['annotations']):
                leg = rte.legs[i]
                leg.annotations = [VFRAnnotation.fromDict(ann, leg) for ann in l]
            rte.set_state(VFRRouteState.ANNOTATIONS)
        # TODO: step 5: tracks
        if state.value>=VFRRouteState.FINALIZED.value:
            rte.tracks = [VFRTrack.fromDict(t, rte) for t in jsonrte['step5']['tracks']]
            rte.set_state(VFRRouteState.FINALIZED)
        # set final state and return
        rte.set_state(state)
        return rte
