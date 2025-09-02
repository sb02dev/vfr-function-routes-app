# coding: utf-8
"""
Calculates VFR routes where the legs are defined as a function
"""
# general packages
import json
from pathlib import Path
from typing import Optional, Union
import textwrap
import datetime
import os
import io
import math

import requests

# projection related packages
from pyproj import Proj, Geod
import numpy as np

# pdf and imaging related packages
import PIL
import matplotlib
matplotlib.use("Agg")
# pylint: disable=wrong-import-position
import matplotlib.pyplot as plt

# document creation related packages
from docx import Document
from docx.shared import Cm

# gpx read and create
import gpxpy

# package imports
from .projutils import (
    PointLonLat, PointXY,
    ExtentLonLat, ExtentXY,
    _calculate_2d_transformation_matrix,
    _apply_transformation_matrix,
    _get_extent_from_points,
    _get_extent_from_extents,
)
from .docxutils import add_formula_par
from .rendering import SimpleRect
from .maps import MapDefinition, MapManager
from .geometry import VFRRouteState, VFRLeg, VFRTrack, VFRPoint, VFRCoordSystem, VFRAnnotation
from .linear_approximation import rdp
from .imageutils import paste_img
# pylint: enable=wrong-import-position


class VFRFunctionRoute:  # pylint: disable=too-many-instance-attributes,disable=too-many-public-methods
    """
    A class that can be used to
      - gradually build up a route
      - add flight tracks
      - generate a map of the route and tracks
      - generate a flight plan document
      - generate a flight plan file for SkyDaemon
    """

    HIGH_DPI = int(os.getenv("HIGH_DPI", "600"))
    LOW_DPI = int(os.getenv("LOW_DPI", "72"))
    DOC_DPI = int(os.getenv("DOC_DPI", "200"))


    @property
    def use_realtime_data(self):
        """A property returning the current value of wheater we should use
        realtime data in our image drawing
        """
        return self._use_realtime_data

    @use_realtime_data.setter
    def use_realtime_data(self, val):
        """A property setter for wheater we should use
        realtime data in our image drawing
        """
        self._use_realtime_data = val
        for l in self.legs:
            for a in l.annotations:
                a.clear_cache()

    @property
    def state(self):
        """A read-only property to access the current state"""
        return self._state

    @property
    def session(self):
        """A read-only property to access the requests.session assigned to this route"""
        return self._session

    @property
    def proj(self):
        """A read-only property to access the projection object"""
        return self._proj

    @property
    def geod(self):
        """A read-only property to access the geod calculation object"""
        return self._geod

    @property
    def matrix_fullmap2map(self):
        """A read-only property to access the transformation matrix"""
        return self._matrix_fullmap2map

    @property
    def matrix_map2fullmap(self):
        """A read-only property to access the inverse transformation matrix"""
        return self._matrix_map2fullmap

    @property
    def matrix_map2cropmap(self):
        """A read-only property to access the transformation matrix"""
        return self._matrix_map2cropmap

    @property
    def matrix_cropmap2map(self):
        """A read-only property to access the inverse transformation matrix"""
        return self._matrix_cropmap2map

    def __init__(self,  # pylint: disable=too-many-arguments,disable=too-many-positional-arguments
                 name: str,
                 mapdef: MapDefinition,
                 speed: float,
                 dof: datetime.datetime,
                 session: Optional[requests.Session] = None,
                 workfolder: Union[str, Path, None] = None,
                 outfolder: Union[str, Path, None] = None,
                 tracksfolder: Union[str, Path, None] = None
                ):
        """
        """
        self._state = VFRRouteState.INITIATED
        self.name = name
        self.map = mapdef
        self.speed = speed
        self.dof = dof
        self.workfolder = workfolder
        self.outfolder = outfolder
        self.tracksfolder = tracksfolder
        self.legs: list[VFRLeg] = []
        self.tracks: list[VFRTrack] = []
        self._session = session if session else requests.Session()
        self.waypoints: list[tuple[str, VFRPoint]] = []
        self.area_of_interest = {
            'top-left': VFRPoint(
                self.map.area["top-left"].lon,
                self.map.area["top-left"].lat,
                VFRCoordSystem.LONLAT, self),
            'bottom-right': VFRPoint(
                self.map.area["bottom-right"].lon,
                self.map.area["bottom-right"].lat,
                VFRCoordSystem.LONLAT, self)
        }
        self._use_realtime_data = False
        self.calc_extents()
        self.calc_transformations()


    def ensure_state(self,
                      required_state: VFRRouteState,
                      ensure_minimum: bool = True,
                      ensure_exactly: bool = False):
        """
        Ensure the object is in the desired state. Otherwise raise an exception.

        Args
            required_state: VFRRouteState
                The state we compare to

            ensure_minimum: bool
                The direction we compare: if True we need to have at least
                the given state, if False we need to have at most the given
                state.

            ensure_exactly: bool
                We have to have exactly the given state (no more, no less).
        """
        if ensure_exactly:
            if self._state.value!=required_state.value:
                raise RuntimeError(
                    "VFRFunctionRoutes object not in required state: " + \
                    f"Current {self._state}, required exact state: {required_state}."
                )
        elif ensure_minimum:
            if self._state.value<required_state.value:
                raise RuntimeError(
                    "VFRFunctionRoutes object not in required state: " + \
                    f"Current {self._state}, required minimum state: {required_state}."
                )
        else:
            if self._state.value>required_state.value:
                raise RuntimeError(
                    "VFRFunctionRoutes object not in required state: " + \
                    f"Current {self._state}, required maximum state: {required_state}."
                )


    def set_state(self, required_state: VFRRouteState):
        """Set the state of the route. It does the neccessary calculations
        for the state transitions

        Args
            required_state: VFRRouteState
                The state we transition to
        """
        if self._state==required_state:
            return
        if self._state.value<required_state.value: # forward stepping
            if self._state == VFRRouteState.INITIATED and required_state.value > self._state.value:
                # INITIADED -> AREAOFINTEREST
                self._state = VFRRouteState.AREAOFINTEREST
                self.calc_extents()
                self.calc_transformations()
            if self._state==VFRRouteState.AREAOFINTEREST and required_state.value>self._state.value:
                # AREAOFINTEREST -> WAYPOINTS
                self._state = VFRRouteState.WAYPOINTS
                self.waypoints_to_legs()
                self.calc_extents()
                self.calc_transformations()
            if self._state == VFRRouteState.WAYPOINTS and required_state.value > self._state.value:
                # WAYPOINTS -> LEGS
                self._state = VFRRouteState.LEGS
                self.legs_to_annotations()
                self.calc_extents()
                self.calc_transformations()
            if self._state == VFRRouteState.LEGS and required_state.value > self._state.value:
                # LEGS -> ANNOTATIONS
                self._state = VFRRouteState.ANNOTATIONS
            if self._state == VFRRouteState.ANNOTATIONS and \
                    required_state.value > self._state.value:
                # ANNOTATIONS -> FINALIZED
                self.finalize()
        else: # backward stepping
            self._state = required_state


    def waypoints_to_legs(self):
        """
        Converts waypoints to legs considering the already existing ones
        """
        for i, wp_start in enumerate(self.waypoints):
            wp_end = self.waypoints[i+1 if i+1<len(self.waypoints) else 0]
                # circle around (last point is the same as first)
            if len(self.legs)>i: # we have a leg at that position
                leg = self.legs[i]
                # so we adjust its endpoints position (not the x value)
                leg.points[0] = (VFRPoint(wp_start[1].lon,
                                          wp_start[1].lat,
                                          VFRCoordSystem.LONLAT, self),
                                 leg.points[0][1])
                leg.points[-1] = (VFRPoint(wp_end[1].lon,
                                           wp_end[1].lat,
                                           VFRCoordSystem.LONLAT, self),
                                  leg.points[-1][1])
                # we adjust the name of the leg
                leg.name = f"{wp_start[0]} -- {wp_end[0]}"
                # we adjust the annotations so we have the first and last match
                self.set_state(VFRRouteState.LEGS)
                    # needed for the annotations but at this point we already are in that state
                if len(leg.annotations)>0:
                    leg.annotations[0].x = leg.points[0][1]
                else:
                    leg.add_annotation('???', leg.points[0][1], (0,0))
                if len(leg.annotations)>1:
                    leg.annotations[-1].x = leg.points[-1][1]
                else:
                    leg.add_annotation('???', leg.points[-1][1], (0,0))
            else: # we don't have a leg yet
                # so we add a new one
                func_range = f"x=0\\textrm{{ at {wp_start[0]}, }}x=1\\textrm{{ at {wp_end[0]}}}"
                self.add_leg(f"{wp_start[0]} -- {wp_end[0]}", f"x^{i+1}",
                             func_range,
                             [
                                 (VFRPoint(wp_start[1].lon,
                                           wp_start[1].lat,
                                           VFRCoordSystem.LONLAT, self),
                                  0),
                                 (VFRPoint(wp_end[1].lon,
                                           wp_end[1].lat,
                                           VFRCoordSystem.LONLAT, self),
                                  1)
                             ])


    def legs_to_annotations(self):
        """
        Transfers information from leg points to annotations (at least a start and an end is needed)
        """
        for leg in self.legs:
            if len(leg.annotations) > 0:
                leg.annotations[0].x = leg.points[0][1]
            else:
                leg.add_annotation('???', leg.points[0][1], (0, 0))
            if len(leg.annotations) > 1:
                leg.annotations[-1].x = leg.points[-1][1]
            else:
                leg.add_annotation('???', leg.points[-1][1], (0, 0))


    def finalize(self):
        """Put the object in the FINALIZED state"""
        self.ensure_state(VFRRouteState.ANNOTATIONS)
        self.calc_extents()
        self.calc_transformations()
        self._state = VFRRouteState.FINALIZED


    def set_area_of_interest(self,
                             top_left_x: float,
                             top_left_y: float,
                             bottom_right_x: float,
                             bottom_right_y: float
                            ) -> None:
        """Change the area of interest rectangle based on
        full map x-y coordinates.
        """
        self.ensure_state(VFRRouteState.INITIATED)
        self.area_of_interest = {
            'top-left': VFRPoint(top_left_x,
                                 top_left_y,
                                 VFRCoordSystem.MAP_XY, self)
                                 .project_point(VFRCoordSystem.LONLAT),
            'bottom-right': VFRPoint(bottom_right_x,
                                     bottom_right_y,
                                     VFRCoordSystem.MAP_XY, self)
                                     .project_point(VFRCoordSystem.LONLAT)
        }

    def set_area_of_interest_lonlat(self,
                                    top_left_lon: float,
                                    top_left_lat: float,
                                    bottom_right_lon: float,
                                    bottom_right_lat: float
                                   ) -> None:
        """Change the area of interest rectangle based on
        world longitude-latitude coordinates.
        """
        self.ensure_state(VFRRouteState.INITIATED)
        self.area_of_interest = {
            'top-left': VFRPoint(top_left_lon, top_left_lat, VFRCoordSystem.LONLAT, self),
            'bottom-right': VFRPoint(bottom_right_lon,
                                     bottom_right_lat,
                                     VFRCoordSystem.LONLAT, self)
        }

    def _get_image_from_figure(self,
                               fig,
                               size: Optional[tuple[float, float]] = None,
                               dpi: Optional[float] = None
                              ) -> io.BytesIO:
        """Private helper function to get a MatPlotLib Figure converted
        to a byte buffer with PNG format image data.
        """
        buf = io.BytesIO()
        if size:
            figsize = fig.get_size_inches()
            dpi = min(size[0] / figsize[0], size[1] / figsize[1])
        fig.savefig(buf, format="png", dpi=dpi, transparent=True)
        buf.seek(0)
        return buf


    def draw_annotations(self):
        """Draw the route with annotations into a MatPlotLib Figure.
        Used for SVG conversion to later serve to the frontend for local drawing.
        """
        fig = plt.figure()
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)

        calc_time = 0
        for l in self.legs:
            calc_time += l.draw(ax, True)

        print(f'calculation time: {calc_time:15,d}')

        return fig


    def draw_tracks(self):
        """Draw the Route without annotations into a MatPlotLib Figure.
        Used for SVG conversion to later serve to the frontend for local drawing.
        """
        fig = plt.figure()
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)

        for l in self.legs:
            l.draw(ax, False)
        for t in self.tracks:
            t.draw(ax)

        return fig


    def add_waypoint(self, name: str, point: VFRPoint):
        """Add a new waypoint to the Route"""
        point.route = self
        self.waypoints.append((name, point.project_point(VFRCoordSystem.LONLAT)))


    def update_waypoints(self, wps: list[dict]):
        """Update the waypoints based on the data received from the frontend."""
        # calculate new waypoints
        self.waypoints = [(
            wp["name"],
                VFRPoint(wp["x"],
                         wp["y"],
                         VFRCoordSystem.MAPCROP_XY, self)
                         .project_point(VFRCoordSystem.LONLAT)
            if 'x' in wp and 'y' in wp else
                VFRPoint(wp["lon"],
                         wp["lat"],
                         VFRCoordSystem.LONLAT, self)
        ) for wp in wps]


    def update_legs(self, legs: list[dict]):
        """Update the legs based on the data received from the frontend."""
        # set legs according to edits
        for i, leg in enumerate(legs):
            curleg = self.legs[i]
            # setup general
            curleg.name = leg["name"]
            # setup function
            curleg.function_range = leg["function_range"]
            latex = leg["function_name"]
            curleg.function_name = latex
            curleg.calc_function()
            # setup constraint points
            lp_start, lp_end = curleg.points[0], curleg.points[-1]
            newpoints: list[tuple[VFRPoint, float]] = []
            for j, pt in enumerate(leg["points"]):
                if j==0:
                    newpoints.append((lp_start[0], pt["func_x"]))
                elif j==len(leg["points"])-1:
                    newpoints.append((lp_end[0], pt["func_x"]))
                else:
                    if 'x' in pt and 'y' in pt:
                        p = VFRPoint(pt["x"],
                                     pt["y"],
                                     VFRCoordSystem.MAPCROP_XY, self, curleg) \
                                     .project_point(VFRCoordSystem.LONLAT)
                    else:
                        p = VFRPoint(pt["lon"],
                                     pt["lat"],
                                     VFRCoordSystem.LONLAT, self, curleg)
                    newpoints.append((
                        p,
                        pt["func_x"]
                    ))
            curleg.points = newpoints
            # recalculate
            curleg.calc_transformations()


    def update_annotations(self, legs: list[dict]):
        """Update the annotations based on the data received from the frontend."""
        for i, l in enumerate(legs):
            if i>len(self.legs)-1:
                break
            if self.legs[i].name!=l['name']:
                print(f"WARNING: leg number {i} name does not match "+
                      f"({self.legs[i].name}!={l['name']})")
            self.legs[i].annotations = [VFRAnnotation(self.legs[i],
                                                      a['name'],
                                                      a['func_x'],
                                                      (a['ofs']['x'], a['ofs']['y']))
                                        for a in l['annotations']
                                       ]


    def add_leg(self,
                name: str,
                function_name: str,
                function_range: str,
                points: list[tuple[VFRPoint, float]]) -> VFRLeg:
        """Initialize and add a leg to the current route"""
        self.ensure_state(VFRRouteState.WAYPOINTS)
        for p, _ in points:
            p.route = self
        newleg = VFRLeg(self, name, function_name, function_range, points)
        self.legs.append(newleg)
        return newleg

    def add_track(self,
                  fname: Union[str, Path],
                  color: str,
                  xmlb: Optional[bytes] = None):
        """Adds a flown track to show on the map.
        
        Args:
            fname: str
                The filename of the track (it will be looked for in the trackfolders folder
                if xmlstring is not given)
            color: str
        """
        self.ensure_state(VFRRouteState.ANNOTATIONS)
        # ensure no name clash
        (ofname, ofext), i = os.path.splitext(fname), 0
        while fname in (t.fname for t in self.tracks):
            fname = f"{ofname}_{i:03d}{ofext}"
            i += 1
        # add the track
        self.tracks.append(VFRTrack(self, fname, color, xmlb=xmlb))
        return self


    def update_tracks(self, tracks):
        """Update the tracks based on the data received from the frontend."""
        newtracks: list[VFRTrack] = []
        for t in self.tracks:
            if t.fname in [nt['name'] for nt in tracks]:
                newtracks.append(t)
                t.color = [nt['color'] for nt in tracks if nt['name']==t.fname][0]
        self.tracks = newtracks


    def calc_transformations(self):
        """Calculates and saves the transformations between coordinate systems.
        
        The following coordinate systems are neccessary:
            1. the function coordinate system
            2. cropped map image coordinates
            3. the map image coordinate system
            4. full-world x-y coordinates
            5. the latitude-longitude coordinates on the earth
            
        It is neccessary to transform between those coordinate systems:
            - 1->2 to draw the route on the image
            - 1->5 to save the route to a flightplan
            - 5->2 to draw tracks on the map image
        Problem:
            2 through 5 is map dependent
            1 is function (leg) dependent
        """
        # to consider:
        #   map false_easting, false_northing values calc/hardcoded
        #   map scale
        # calculate transformations for LONLAT<->FULL_WORLD_XY
        self._proj_str = "+proj=lcc +lon_0=-90 +lat_1=46 +lat_2=48 +ellps=WGS84"
        self._proj = Proj(self._proj_str)
        self._geod = Geod(self._proj_str)
        #print(self._proj(16,48.5))
        #print(self._proj(17,47.5))
        # calculate transformations for FULL_WORLD_XY<->MAP_XY
        p = [PointXY(x, y) for x, y in [self._proj(ll.lon, ll.lat)
             for ll in self.map.points.keys()]] # convert to fullworld map coord
        pp = [PointXY(pxy.x/72*self.LOW_DPI, pxy.y/72*self.LOW_DPI)
              for pxy in self.map.points.values()]
                # must scale it to LOW_DPI from default pdf metric of 72
        self._matrix_fullmap2map = _calculate_2d_transformation_matrix(p, pp)
            # calc matrix from fullworld map coord to map coord
        self._matrix_map2fullmap = np.linalg.inv(self._matrix_fullmap2map)
        #print("TEST lonlat to mapxy:")
        #for i, (lonlat, xy) in enumerate(self.PDF_IN_WORLD_XY.items()):
        #    print(f"  {lonlat} -> {p[i]} -> {xy} vs "+
        #          f"{_apply_transformation_matrix(p[i], self._matrix_fullmap2map)}")
        # calculate transformations for MAP_XY<->MAPCROP_XY
        p = self.get_mapxyextent()
        pp = [
            (0, 0),
            (0, (p.maxy-p.miny)*self.HIGH_DPI/self.LOW_DPI),
            ((p.maxx-p.minx)*self.HIGH_DPI/self.LOW_DPI, 0),
            ((p.maxx-p.minx)*self.HIGH_DPI/self.LOW_DPI,
             (p.maxy-p.miny)*self.HIGH_DPI/self.LOW_DPI),
        ]  # also scale it up!
        p = [
            PointXY(p.minx, p.miny),
            PointXY(p.minx, p.maxy),
            PointXY(p.maxx, p.miny),
            PointXY(p.maxx, p.maxy),
        ]
        self._matrix_map2cropmap = _calculate_2d_transformation_matrix(p, pp)
        self._matrix_cropmap2map = np.linalg.inv(self._matrix_map2cropmap)
        #print("TEST mapxy to cropmapxy:")
        #for i, (mapxy, cropmapxy) in enumerate(zip(p, pp)):
        #    print(f"  {mapxy} -> {cropmapxy} vs "+
        #          f"{_apply_transformation_matrix(mapxy, self._matrix_map2cropmap)}")
        # calculate transformations for each leg
        for leg in self.legs:
            leg.calc_transformations()


    def calc_extents(self, margin_x: float = .2, margin_y: Optional[float] = None):
        """Calculates and saves the extents of the neccessary map.
        
        Calculates the top-left and bottom-right coordinates of the map
        that will contain all points of the route. It also adds a margin
        in both directions given in the parameters as a percentage.

        If there are no legs nor tracks, than it defaults to an area around Budapest.
        
        Args:
            margin_x:
                Horizontal additional margin. Given in percentage of the
                distance of minimum and maximum lateral coordinates.
                Default is 10%
            margin_y:
                Optional vertical margin. Given in percentage of the
                distance of minimum and maximum longitudinal coordinates.
                Default is equal to horizontal margin.
        """
        lat0, lat1, lon0, lon1 = (self.area_of_interest["top-left"].lat,
                                    self.area_of_interest["bottom-right"].lat,
                                    self.area_of_interest["top-left"].lon,
                                    self.area_of_interest["bottom-right"].lon)
        area_of_interest = ExtentLonLat(
            min(lon0, lon1),
            min(lat0, lat1),
            max(lon0, lon1),
            max(lat0, lat1)
        )
        if self._state in [VFRRouteState.INITIATED,
                           VFRRouteState.AREAOFINTEREST,
                           VFRRouteState.WAYPOINTS]:
            # at this stage we only have area of interest
            # (which is always set, at least to a default)
            self.extent = area_of_interest
        elif self._state in [VFRRouteState.LEGS,
                             VFRRouteState.ANNOTATIONS,
                             VFRRouteState.FINALIZED]:
            # get the automatic bounding box
            if len(self.legs)==0 and len(self.tracks)==0:
                # no legs, no tracks fall back to waypoints extent
                if len(self.waypoints)==0:
                    self.extent = area_of_interest
                    return

                extent = _get_extent_from_points([PointLonLat(p.lon, p.lat)
                                                  for name, p in self.waypoints])
            else:
                # we have at least one leg or track, use those (leg points contain waypoints)
                extent = _get_extent_from_extents(
                        [l.get_extent() for l in self.legs] +
                        [t.get_extent() for t in self.tracks]
                    )
            # add margins
            if margin_y is None:
                margin_y = margin_x
            extent_with_margins = ExtentLonLat(
                    extent.minlon - (extent.maxlon-extent.minlon)*margin_x,
                    extent.minlat - (extent.maxlat-extent.minlat)*margin_y,
                    extent.maxlon + (extent.maxlon-extent.minlon)*margin_x,
                    extent.maxlat + (extent.maxlat-extent.minlat)*margin_y
                )
            # get the bounding box of the automatic and the manually defined
            # (i.e. only increase manually given box)
            extent_or_aoi = ExtentLonLat(
                min(extent_with_margins.minlon, area_of_interest.minlon),
                min(extent_with_margins.minlat, area_of_interest.minlat),
                max(extent_with_margins.maxlon, area_of_interest.maxlon),
                max(extent_with_margins.maxlat, area_of_interest.maxlat)
            )
            self.extent = extent_or_aoi


    def get_mapxyextent(self) -> ExtentXY:
        """Get the extent of the Route.
        It is calculated based on the points of the functions.
        """
        p = [
            PointLonLat(self.extent.minlon, self.extent.minlat), # minlon-minlat -> leftbottom
            PointLonLat(self.extent.minlon, self.extent.maxlat), # minlon-maxlat -> lefttop
            PointLonLat(self.extent.maxlon, self.extent.minlat), # maxlon-minlat -> rightbottom
            PointLonLat(self.extent.maxlon, self.extent.maxlat), # maxlon-maxlat -> righttop
        ]
        p = [self._proj(ll.lon, ll.lat) for ll in p] # projected to FULL_WORLD_XY
        p = [_apply_transformation_matrix(pp, self._matrix_fullmap2map)
             for pp in p] # projected to MAP_XY (LOW_DPI)
        p = [PointXY(pp[0], pp[1]) for pp in p]
        return ExtentXY(
            min(pp.x for pp in p),
            min(pp.y for pp in p),
            max(pp.x for pp in p),
            max(pp.y for pp in p)
        )


    def draw_map(self, use_realtime: Optional[bool] = None):  # pylint: disable=too-many-locals
        """Draws a matplotlib based map of the defined route.
        
        Args:
        
        Returns:
            The matplotlib axes of the final plot.
        """
        self.ensure_state(VFRRouteState.FINALIZED)

        # save the realtime data usage status
        set_rtd, old_rtd = False, self.use_realtime_data
        if use_realtime is not None:
            if use_realtime!=self.use_realtime_data:
                old_rtd, self.use_realtime_data, set_rtd = self.use_realtime_data, True, True

        # save the original clear function (to be restored in finally)
        from matplotlib.backends import backend_agg # pylint: disable=import-outside-toplevel
        orig_clear = backend_agg.RendererAgg.clear

        image_size = PointXY(800, 600)

        try:
            # setup clear function to draw background
            tiles = self.map.get_tilerenderer(int(os.getenv('DOC_DPI', '200')))
            if tiles is not None:
                tile_list, crop, image_size, tile_range = \
                    tiles.get_tile_list_for_area(self.calc_basemap_clip())
                def custom_background(renderer):
                    arr = np.asarray(renderer.buffer_rgba())
                    for p in tile_list:
                        tile = np.asarray(tiles.get_tile(p.x, p.y)[1])
                        # we need to shift the images, cropping not needed (its outside anyway)
                        x = int((p.x - tile_range[0])*tiles.tile_size[0] - crop.p0.x)
                        y = int((p.y - tile_range[2])*tiles.tile_size[1] - crop.p0.y)
                        paste_img(arr, tile, x, y)
                backend_agg.RendererAgg.clear = custom_background # type: ignore

            # initialize map
            from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas  # pylint: disable=import-outside-toplevel
            fig = plt.figure()
            canvas = FigureCanvas(fig)
            ax = plt.Axes(fig, [0., 0., 1., 1.])
            ax.set_axis_off()
            fig.add_axes(ax)

            # draw the map parts
            for l in self.legs:
                l.draw(ax)
            for t in self.tracks:
                t.draw(ax)

            # render the overlay
            fig.patch.set_alpha(0.0)      # transparent background instead of white # type: ignore
            ax.patch.set_alpha(0.0)       # same for axes # type: ignore
            fig.set_dpi(self.DOC_DPI)
            fig.set_size_inches(image_size.x/self.DOC_DPI,  image_size.y/self.DOC_DPI)
            ax.set_xlim(0, image_size.x/self.DOC_DPI*self.HIGH_DPI)
            ax.set_ylim(image_size.y/self.DOC_DPI*self.HIGH_DPI, 0)
            canvas.draw()
            img_buf = canvas.buffer_rgba()
            img = PIL.Image.frombuffer("RGBA", # type: ignore
                                       (int(image_size.x), int(image_size.y)),
                                       img_buf,
                                       "raw",
                                       "RGBA", 0, 1)
            plt.close(fig)

        finally:
            # restore realtime wind state
            if set_rtd:
                self.use_realtime_data = old_rtd
            # restore matplotlib default
            backend_agg.RendererAgg.clear = orig_clear

        # return the composited
        buf = io.BytesIO()
        img.save(buf, 'png')
        buf.seek(0)
        return buf.getvalue()


    def calc_basemap_clip(self) -> SimpleRect:
        """Calculates the rectangle of the desired area on the map in
        PDF coordinate system.
        """
        # calc clip coordinates from the appropriate lon-lat corners
        lat0, lat1, lon0, lon1 = self.extent.minlat, self.extent.maxlat, \
                                 self.extent.minlon, self.extent.maxlon
        # adjust for non-rectangle because of projection type
        corners_lonlat = [
            VFRPoint(lon0, lat0, route=self),
            VFRPoint(lon1, lat1, route=self),
            VFRPoint(lon1, lat0, route=self),
            VFRPoint(lon0, lat1, route=self)
        ]
        corners_map = [p.project_point(to_system=VFRCoordSystem.MAP_XY) for p in corners_lonlat]
        x0 = min(p.x for p in corners_map)
        y0 = min(p.y for p in corners_map)
        x1 = max(p.x for p in corners_map)
        y1 = max(p.y for p in corners_map)
        # the order of them is important
        if y1<y0:
            y0, y1 = y1, y0
        # this is in LOW_DPI => convert it back to PDF coordinates
        ((x0, y0), (x1, y1)) = ((x0/self.LOW_DPI*72, y0/self.LOW_DPI*72),
                                (x1/self.LOW_DPI*72, y1/self.LOW_DPI*72))
        ((xm, ym), (_, _)) = self.map.margins
        ((x0, y0), (x1, y1)) = ((xm+x0, ym+y0), (xm+x1, ym+y1))
        return SimpleRect(PointXY(x0, y0), PointXY(x1, y1))

    def create_doc(self,  # pylint: disable=too-many-locals
                   save: bool = True
                  ) -> Union[io.BytesIO, None]:
        """
        Generates a report of the route as a Word document
        Argument: a list of a tuple
        - first element is a tuple of the leg data
            - 1st element is the name (section header)
            - 2nd element is the function with the bounds used as a text
            - 3rd element is the definition of x points to real life as a text
        - second element is a list of segment data, which is a tuple
            - 1st element is the name
            - 2nd element is the heading at that point
            - 3rd element is the length of the curve segment in meters
            - 4th element is the time used for that curve segment
        """
        self.ensure_state(VFRRouteState.FINALIZED)
        # draw map if we don't have it yet and save the image
        old_rtd, set_rtd = False, False
        if not self.use_realtime_data:
            old_rtd, self.use_realtime_data, set_rtd = self.use_realtime_data, True, True
        image = self.draw_map()
        imgname = os.path.join(self.outfolder if self.outfolder is not None
                               else '', self.name+'.png')
        with open(imgname, "wb") as f:
            f.write(image)

        # header and image
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        doc.add_heading('Route Plan', 0)
        doc.add_paragraph(f"Planned speed: {self.speed} KIAS.")
        doc.add_paragraph(f"Wind forecast for {self.dof:%Y-%m-%d %H:%M %Z}.")
        doc.add_picture(f'{imgname}', width=Cm(19.00))
        doc.add_page_break()

        # legs
        totdist, tottime, tottimewc = 0, 0, 0
        for leg in self.legs:
            # leg heading and definiton
            doc.add_heading(leg.name, level=1)
            doc.add_heading("Definition", level=2)
            add_formula_par(doc, leg.function_name, style="List Bullet")
            add_formula_par(doc, leg.function_range, style="List Bullet")
            doc.add_heading("Segments", level=2)

            # leg table header
            tab = doc.add_table(rows=1, cols=8)
            tab.autofit = True
            tab.allow_autofit = True  # type: ignore
            tab.style = "Colorful Shading Accent 1"
            for i, hdr in enumerate(["Name",
                                     "Hdg",
                                     "Mag",
                                     "WCA",
                                     "Length",
                                     "Time",
                                     "Tme(WC)",
                                     "Wind"]):
                tab.rows[0].cells[i].text = hdr

            # leg table rows (per annotations)
            for ann in leg.annotations:
                curdist, curtime, curtimewc = self.add_annotation_to_doc(tab, ann)
                totdist += curdist
                tottime += curtime
                tottimewc += curtimewc

        # total distance/time
        doc.add_paragraph(f"Total: {totdist/1852:.0f} NM, " +
                          f"{math.floor(tottime):2d}: \
                            {math.floor((tottime-math.floor(tottime))*60):02d} / " +
                          f"{math.floor(tottimewc):2d}: \
                            {math.floor((tottimewc-math.floor(tottimewc))*60):02d}"
                         )


        # save it
        if set_rtd:
            self.use_realtime_data = old_rtd
        if save:
            docname = os.path.join(self.outfolder if self.outfolder is not None
                                   else '', self.name+'.docx')
            doc.save(docname)
            return None

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def add_annotation_to_doc(self, tab, ann: VFRAnnotation):
        """Adds an annotation to a table in a document.
        
        Args
            tab: the table in the word document
            ann: the annotation to add

        Returns:
            tuple: distance, time, wind corrected time of the current
                   annotation segment of the route
        """
        seglen = ann.seglen
        segtime = ann.segtime
        segtime_wind = sum(ann.times_withwind)
        seghdgs = ann.headings
        seghdg = seghdgs[-1]
        wind_corrs = ann.wind_corrections(headings=seghdgs)
        wind_corr = wind_corrs[-1] if wind_corrs else None
        magdev = ann.magnetic_deviation(self.dof)


        row_cells = tab.add_row().cells

        row_cells[0].text = ann.name
                #row_cells[0].width = Cm(2)

        row_cells[1].text = f"{seghdg:5.0f}\N{DEGREE SIGN}"
        row_cells[2].text = f"{magdev:3.0f}\N{DEGREE SIGN}"
        row_cells[3].text = f"{wind_corr:3.0f}\N{DEGREE SIGN}"
        row_cells[4].text = \
                    f"{seglen/1852 if seglen is not None else '-' \
                       :{'' if seglen is None else '.1f'}}NM"
        row_cells[5].text = f"{math.floor(segtime) if segtime is not None \
                                       else '-':{'' if segtime is None else '2d'}}" + \
                                    f":{math.floor((segtime-math.floor(segtime))*60) \
                                        if segtime is not None else '--' \
                                        :{'' if segtime is None else '02d'}}"
        row_cells[6].text = f"{math.floor(segtime_wind) if segtime_wind is not None \
                                       else '-':{'' if segtime_wind is None else '2d'}}"+\
                                    f":{math.floor((segtime_wind-math.floor(segtime_wind))*60) \
                                        if segtime_wind is not None else '--' \
                                        :{'' if segtime_wind is None else '02d'}}"
        row_cells[7].text = f"{ann.wind_dir:3d}\N{DEGREE SIGN} {ann.wind_speed:.0f}kts"

        curdist = seglen if seglen is not None else 0
        curtime = segtime if segtime is not None else 0
        curtimewc = segtime_wind if segtime_wind is not None else 0
        return curdist, curtime, curtimewc


    def save_plan(self):
        """Get an XML string of a GPX format of the Route.
        This is importable into SkyDemon.
        The functions are approximated by straight lines
        (otherwise SkyDemon really slows down).
        """
        self.ensure_state(VFRRouteState.FINALIZED)
        gpx = gpxpy.gpx.GPX()  # type: ignore
        gpx.name = "Elmebeteg VFR útvonal"
        gpx.time = datetime.datetime.now()
        rte = gpxpy.gpx.GPXRoute(name="Elmebeteg VFR útvonal")  # type: ignore
        for leg in self.legs:
            x = np.linspace(min(x for p, x in leg.points),
                            max(x for p, x in leg.points),
                            500
                            )
            breakpoints = rdp(np.column_stack((x, np.array([leg.function(xx) for xx in x]))), 0.025)
            psrc = [VFRPoint(x,
                             y,
                             VFRCoordSystem.FUNCTION, self, leg)
                    for x, y in breakpoints]
            ps = [p.project_point(VFRCoordSystem.LONLAT) for p in psrc]
            pt = [gpxpy.gpx.GPXRoutePoint(p.lat,  # type: ignore
                                          p.lon,
                                          name=leg.name if i==0 else None)
                  for i, p in enumerate(ps)]
            rte.points.extend(pt)
        gpx.routes.append(rte)
        return gpx.to_xml()


    def __repr__(self):
        """String representation of the object"""
        s = f"{type(self).__name__}({self.name})\n"
        for l in self.legs:
            s += textwrap.indent(repr(l), "  ")
        return s


    def to_dict(self):
        """Converts the object to serializable dictionary."""
        # initiate json object with basic info
        jsonrte = {
            'name': self.name,
            'mapname': self.map.name,
            'speed': self.speed,
            'dof': self.dof.isoformat(),
            'state': self._state.name
        }
        # step 1: area of interest
        #if self._state.value>=VFRRouteState.AREAOFINTEREST.value:
        jsonrte['step1'] = { 'area_of_interest': {
            'top-left': self.area_of_interest['top-left'].to_dict(),
            'bottom-right': self.area_of_interest['bottom-right'].to_dict(),
        }}
        # step 2: waypoints
        #if self._state.value>=VFRRouteState.WAYPOINTS.value:
        jsonrte['step2'] = { 'waypoints': [(wp[0], wp[1].to_dict()) for wp in self.waypoints] }
        # step 3: legs
        #if self._state.value>=VFRRouteState.LEGS.value:
        jsonrte['step3'] = { 'legs': [leg.to_dict() for leg in self.legs] }
        # step 4: annotation points
        #if self._state.value>=VFRRouteState.ANNOTATIONS.value:
        jsonrte['step4'] = {
            'annotations': [[ann.to_dict()
                             for ann in leg.annotations]
                            for leg in self.legs]
            }
        # step 5: tracks
        #if self._state.value>=VFRRouteState.FINALIZED.value:
        jsonrte['step5'] = { 'tracks': [t.to_dict() for t in self.tracks]}
        # return
        return jsonrte

    def to_json(self):
        """Serializes the object to JSON string."""
        return json.dumps(self.to_dict(), indent=2)


    @classmethod
    def from_json(cls, jsonstring: str,
                 session: Optional[requests.Session] = None,
                 workfolder: Union[str, Path, None] = None,
                 outfolder: Union[str, Path, None] = None,
                 tracksfolder: Union[str, Path, None] = None):
        """Deserializes the object from a JSON string."""
        # decode json
        jsonrte = json.loads(jsonstring)
        # load it
        return VFRFunctionRoute.from_dict(jsonrte, session, workfolder, outfolder, tracksfolder)

    @classmethod
    def from_dict(cls, jsonrte: dict,
                 session: Optional[requests.Session] = None,
                 workfolder: Union[str, Path, None] = None,
                 outfolder: Union[str, Path, None] = None,
                 tracksfolder: Union[str, Path, None] = None):
        """Deserializes the object from a dictionary."""
        if MapManager.instance() is None:
            raise ValueError('There is no MapManager initialized')
        # initiate with basic info
        rte = VFRFunctionRoute(jsonrte['name'],
                               MapManager.instance().maps.get(jsonrte['mapname']), # type: ignore
                               jsonrte['speed'],
                               datetime.datetime.fromisoformat(jsonrte['dof']),
                               session, workfolder, outfolder, tracksfolder)
        state = VFRRouteState[jsonrte['state']]
        # step 1: area of interest
        #if state.value>=VFRRouteState.AREAOFINTEREST.value:
        rte.area_of_interest = {
            'top-left': VFRPoint.from_dict(
                jsonrte['step1']['area_of_interest']['top-left'],
                rte),
            'bottom-right': VFRPoint.from_dict(
                jsonrte['step1']['area_of_interest']['bottom-right'],
                rte),
        }
        #rte.set_state(VFRRouteState.AREAOFINTEREST)
        # step 2: waypoints
        #if state.value>=VFRRouteState.WAYPOINTS.value:
        rte.waypoints = [(name, VFRPoint.from_dict(p, rte))
                         for name, p in jsonrte['step2']['waypoints']]
        #rte.set_state(VFRRouteState.WAYPOINTS)
        # step 3: legs
        #if state.value>=VFRRouteState.LEGS.value:
        rte.legs = [VFRLeg.from_dict(leg, rte)
                    for leg in jsonrte['step3']['legs']]
        #rte.set_state(VFRRouteState.LEGS)
        # step 4: annotation points
        #if state.value>=VFRRouteState.ANNOTATIONS.value:
        for i, l in enumerate(jsonrte['step4']['annotations']):
            leg = rte.legs[i]
            leg.annotations = [VFRAnnotation.from_dict(ann, leg)
                               for ann in l]
        #rte.set_state(VFRRouteState.ANNOTATIONS)
        # step 5: tracks
        #if state.value>=VFRRouteState.FINALIZED.value:
        rte.tracks = [VFRTrack.from_dict(t, rte)
                      for t in jsonrte['step5']['tracks']]
        #rte.set_state(VFRRouteState.FINALIZED)
        # set final state and return
        rte.set_state(state)
        return rte
